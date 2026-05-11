#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import zipfile
from copy import deepcopy
from datetime import date
from html import escape
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parent.parent


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def text(value: Any, default: str = "-") -> str:
    if value is None:
        return default
    if isinstance(value, dict):
        for key in ("zh", "text", "name", "label", "en"):
            if value.get(key):
                return str(value[key]).strip()
        return default
    if isinstance(value, list):
        items = [text(item, "") for item in value]
        items = [item for item in items if item]
        return "；".join(items) if items else default
    result = str(value).strip()
    return result if result else default


def as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def number(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def fmt_num(value: Any) -> str:
    n = number(value)
    if abs(n - round(n)) < 1e-9:
        return str(int(round(n)))
    return f"{n:.1f}"


def fmt_pct(value: Any) -> str:
    return f"{number(value) * 100:.0f}%"


def pipe_escape(value: Any) -> str:
    return text(value).replace("|", "\\|").replace("\n", "<br>")


def h(value: Any) -> str:
    return escape(text(value))


def index_by_id(items: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    return {str(item.get("id")): item for item in items if item.get("id")}


def label(ref: Any, action_map: dict[str, dict[str, Any]], default: str | None = None) -> str:
    ref_text = text(ref, "")
    if ref_text in action_map:
        return text(action_map[ref_text].get("name"), ref_text)
    for item in action_map.values():
        if ref_text == text(item.get("name"), ""):
            return text(item.get("name"), ref_text)
    return default or ref_text or "-"


def same_ref(ref: Any, item: dict[str, Any]) -> bool:
    ref_text = text(ref, "")
    return ref_text in {str(item.get("id", "")), text(item.get("name"), "")}


def merge_update(base: dict[str, Any], update: dict[str, Any]) -> dict[str, Any]:
    merged = deepcopy(base)
    append_keys = [
        "players",
        "our_actions",
        "opponent_actions",
        "reaction_estimates",
        "payoff_matrix",
        "commitment_tests",
        "signals",
        "rounds",
        "scenario_triggers",
        "warnings",
    ]
    for key in append_keys:
        if update.get(key):
            merged.setdefault(key, [])
            merged[key].extend(as_list(update[key]))
    for key in ("case_context", "equilibrium", "repeated_game", "recommendation"):
        if isinstance(update.get(key), dict):
            merged.setdefault(key, {})
            merged[key].update(update[key])
    merged.setdefault("applied_updates", [])
    merged["applied_updates"].append(update)
    return merged


def commitment_score(item: dict[str, Any]) -> tuple[float, str]:
    components = [
        number(item.get("costliness"), 0.0),
        number(item.get("observability"), 0.0),
        1.0 - number(item.get("reversibility"), 1.0),
        number(item.get("incentive_fit"), 0.0),
        number(item.get("capability_backing"), 0.0),
    ]
    present = [max(0.0, min(1.0, value)) for value in components]
    score = sum(present) / len(present) if present else 0.0
    if score >= 0.75:
        level = "可信"
    elif score >= 0.45:
        level = "部分可信"
    else:
        level = "弱承诺/更像口头信号"
    return score, level


def enrich_commitments(request: dict[str, Any]) -> list[dict[str, Any]]:
    actions = index_by_id(request.get("our_actions", []))
    enriched = []
    for item in request.get("commitment_tests", []):
        score, level = commitment_score(item)
        enriched.append(
            {
                **item,
                "action_name": label(item.get("action_id"), actions),
                "score": score,
                "level": level,
            }
        )
    return enriched


def payoff_rows_for_action(request: dict[str, Any], our_action: dict[str, Any]) -> list[dict[str, Any]]:
    return [row for row in request.get("payoff_matrix", []) if same_ref(row.get("our_action"), our_action)]


def matching_payoff(row_action: Any, opponent_action: Any) -> bool:
    row = text(row_action, "")
    candidate = text(opponent_action, "")
    return bool(row and candidate and row == candidate)


def expected_payoffs(request: dict[str, Any]) -> list[dict[str, Any]]:
    our_actions = request.get("our_actions", [])
    opponent_actions = index_by_id(request.get("opponent_actions", []))
    results = []
    for action in our_actions:
        action_rows = payoff_rows_for_action(request, action)
        weighted = 0.0
        weight = 0.0
        detail = []
        for estimate in request.get("reaction_estimates", []):
            if not same_ref(estimate.get("if_we"), action):
                continue
            for likely in estimate.get("likely_actions", []):
                p = number(likely.get("probability"), 0.0)
                candidate = text(likely.get("action"), "")
                row = next((r for r in action_rows if matching_payoff(r.get("opponent_action"), candidate)), None)
                if row is None and candidate in opponent_actions:
                    row = next((r for r in action_rows if matching_payoff(r.get("opponent_action"), opponent_actions[candidate].get("name"))), None)
                if row is None:
                    continue
                payoff = number(row.get("our_payoff"), 0.0)
                weighted += p * payoff
                weight += p
                detail.append(
                    {
                        "opponent_action": label(candidate, opponent_actions, candidate),
                        "probability": p,
                        "our_payoff": payoff,
                        "rationale": text(likely.get("rationale")),
                    }
                )
        if weight > 0:
            expected = weighted / weight
            method = "reaction-weighted"
        elif action_rows:
            expected = sum(number(row.get("our_payoff"), 0.0) for row in action_rows) / len(action_rows)
            method = "matrix-average"
        else:
            expected = 0.0
            method = "missing-payoff"
        results.append(
            {
                "action_id": action.get("id"),
                "action_name": text(action.get("name")),
                "expected_payoff": expected,
                "method": method,
                "details": detail,
            }
        )
    return sorted(results, key=lambda item: item["expected_payoff"], reverse=True)


def top_reactions(request: dict[str, Any], action_id: str | None) -> list[dict[str, Any]]:
    if not action_id:
        return []
    opponent_actions = index_by_id(request.get("opponent_actions", []))
    for estimate in request.get("reaction_estimates", []):
        if text(estimate.get("if_we"), "") != action_id:
            continue
        likely = sorted(estimate.get("likely_actions", []), key=lambda item: number(item.get("probability"), 0.0), reverse=True)
        return [
            {
                "opponent": text(estimate.get("opponent")),
                "action": label(item.get("action"), opponent_actions),
                "probability": number(item.get("probability"), 0.0),
                "rationale": text(item.get("rationale")),
            }
            for item in likely
        ]
    return []


def infer_framework_selection(request: dict[str, Any]) -> dict[str, Any]:
    raw = json.dumps(request, ensure_ascii=False).lower()
    primary = "标准型 payoff 矩阵博弈"
    lenses: list[str] = ["纳什均衡检查：确认预测结果是否对各方都是最佳回应"]
    excluded = []
    logic = "根据玩家、策略、payoff、行动顺序和信息结构选择最轻的可解释框架。"

    if any(term in raw for term in ["price", "pricing", "价格", "降价", "折扣", "price war"]):
        primary = "Bertrand 价格竞争"
        lenses.extend(
            [
                "囚徒困境：检查双方是否会被拖入互相降价的低收益均衡",
                "重复博弈：检查短期降价是否破坏长期价格纪律",
            ]
        )
        logic = "价格是核心策略变量，因此先用 Bertrand 竞争解释跟随降价，再用囚徒困境和重复博弈检查价格战风险。"
    if any(term in raw for term in ["channel", "渠道", "partner", "联盟", "生态"]):
        lenses.extend(["联盟/渠道博弈：检查渠道或伙伴是否会因分成、风险和锁定而改变阵营"])
    if any(term in raw for term in ["signal", "信号", "free", "免费", "announcement", "宣布", "roadmap"]):
        lenses.extend(["信号博弈：检查对手或我们的公开动作是否成本高、可验证、难伪造"])
    if any(term in raw for term in ["commitment", "承诺", "exclusive", "独家", "长期"]):
        lenses.extend(["可信承诺：检查动作是否成本高、可观察、难撤回、时间一致"])
    if any(term in raw for term in ["bid", "auction", "m&a", "并购", "竞价", "tender"]):
        primary = "拍卖/竞价博弈"
        lenses.extend(["赢家诅咒：检查中标是否意味着过度乐观估值", "谈判博弈：检查出价后条款和退出选项"])
        logic = "稀缺资产和多方出价是核心结构，因此优先使用拍卖/竞价博弈。"
    if any(term in raw for term in ["financing", "融资", "terms", "估值", "valuation", "batna"]):
        primary = "谈判博弈"
        lenses.extend(["外部选择：检查 BATNA 和保留价值", "信号博弈：检查融资信心与稀缺性信号"])
        logic = "条款和剩余价值分配是核心问题，因此优先使用谈判博弈。"
    if any(term in raw for term in ["entry", "进入", "retaliation", "威慑"]):
        lenses.extend(["进入威慑：检查 incumbent 的威胁是否在进入后仍然理性", "子博弈精炼：过滤不可信威胁"])
    if "zero" not in raw and "零和" not in raw:
        excluded.append(
            {
                "name": "纯零和博弈",
                "reason": "当前输入未证明总收益固定；如果存在渠道、客户体验、合作或差异化空间，应按混合动机博弈处理。",
            }
        )
    return {
        "primary": primary,
        "secondary_lenses": list(dict.fromkeys(lenses)),
        "excluded_frameworks": excluded,
        "combination_logic": logic,
    }


def readiness_status(score: float) -> tuple[str, str]:
    if score >= 0.75:
        return "ready", "关键玩家、策略、payoff、反应和承诺信息已经足够支持正式策略建议。"
    if score >= 0.45:
        return "nearly-ready", "当前模型可用于方向性决策，但仍有一两个信息缺口会影响执行强度。"
    return "needs-more-info", "当前模型仍缺少关键玩家、payoff、反应或承诺证据，适合先做低成本验证。"


def build_sensitivity(request: dict[str, Any], payoffs: list[dict[str, Any]], commitments: list[dict[str, Any]], primary_id: str) -> dict[str, Any]:
    if request.get("sensitivity"):
        return request["sensitivity"]
    best = payoffs[0] if payoffs else {}
    runner_up = payoffs[1] if len(payoffs) > 1 else {}
    gap = number(best.get("expected_payoff"), 0.0) - number(runner_up.get("expected_payoff"), 0.0)
    primary_commitment = next((item for item in commitments if text(item.get("action_id"), "") == primary_id), None)
    commitment_score_value = number(primary_commitment.get("score"), 0.0) if primary_commitment else 0.0
    if gap >= 15 and commitment_score_value >= 0.75:
        stability = "stable"
    elif gap >= 5 or commitment_score_value >= 0.6:
        stability = "mixed"
    else:
        stability = "fragile"
    if gap < 5:
        dangerous = "主推荐与备选方案的预期收益差距较小，少量 payoff 调整就可能改变排序。"
    elif commitment_score_value < 0.75:
        dangerous = "主推荐依赖承诺可信度；如果渠道或对手认为该承诺可轻易撤回，推荐会变弱。"
    else:
        dangerous = "最大风险是假设竞品不会把价格战升级到长期补贴或更强渠道攻击。"
    return {
        "stability": stability,
        "payoff_gap": gap,
        "base_recommendation": text(best.get("action_name"), "-"),
        "runner_up": text(runner_up.get("action_name"), "无"),
        "most_dangerous_assumption": dangerous,
        "stress_tests": [
            {
                "name": "竞品更激进",
                "assumption_shift": "将头部竞品跟随降价或攻击渠道的倾向上调。",
                "expected_effect": "降价选项进一步变差，渠道绑定需要更强返点和交付资源支撑。",
                "recommendation_impact": "通常不改变主推荐，但会提高执行成本。",
            },
            {
                "name": "渠道收益折扣",
                "assumption_shift": "将渠道绑定带来的渠道收益和我方收益下调。",
                "expected_effect": "捆绑渠道与高端差异化的差距缩小。",
                "recommendation_impact": "若渠道拒绝年度合作包，应提高高端差异化权重。",
            },
            {
                "name": "承诺可信度下降",
                "assumption_shift": "渠道认为年度合作包只是短期销售动作。",
                "expected_effect": "渠道绑定的可信承诺分下降，竞品攻击渠道的胜率上升。",
                "recommendation_impact": "需要先补强可观察投入，再正式推进绑定。",
            },
        ],
    }


def build_strategy_readiness(request: dict[str, Any], payoffs: list[dict[str, Any]], commitments: list[dict[str, Any]], framework_selection: dict[str, Any], sensitivity: dict[str, Any]) -> dict[str, Any]:
    if request.get("strategy_readiness"):
        return request["strategy_readiness"]
    checks = {
        "players": 1.0 if len(request.get("players", [])) >= 2 else 0.3,
        "strategies": 1.0 if request.get("our_actions") and request.get("opponent_actions") else 0.35,
        "payoffs": 1.0 if request.get("payoff_matrix") and all(row.get("basis") for row in request.get("payoff_matrix", [])) else 0.45,
        "reactions": 1.0 if any(item.get("likely_actions") for item in request.get("reaction_estimates", [])) else 0.35,
        "commitments": min(1.0, max([number(item.get("score"), 0.0) for item in commitments] or [0.0]) + 0.15),
        "framework": 1.0 if framework_selection.get("primary") and framework_selection.get("secondary_lenses") else 0.4,
        "sensitivity": {"stable": 1.0, "mixed": 0.7, "fragile": 0.35}.get(text(sensitivity.get("stability"), ""), 0.5),
        "updateability": 1.0 if request.get("scenario_triggers") and request.get("rounds") else 0.55,
    }
    weights = {
        "players": 0.12,
        "strategies": 0.12,
        "payoffs": 0.16,
        "reactions": 0.16,
        "commitments": 0.14,
        "framework": 0.1,
        "sensitivity": 0.12,
        "updateability": 0.08,
    }
    score = sum(checks[key] * weights[key] for key in weights)
    if sensitivity.get("stability") == "mixed":
        score = min(score, 0.74)
    elif sensitivity.get("stability") == "fragile":
        score = min(score, 0.44)
    gaps = []
    if checks["payoffs"] < 0.8:
        gaps.append("补充真实毛利、渠道转化、客户留存或竞品成本数据，替换 assumed payoff。")
    if checks["reactions"] < 0.8:
        gaps.append("补齐竞品跟随降价、渠道攻击、免费版本推进的证据和概率。")
    if checks["commitments"] < 0.75:
        gaps.append("提高主推荐动作的可观察投入、合同约束或撤回成本。")
    if sensitivity.get("stability") != "stable":
        gaps.append("围绕最危险假设做一次小规模验证或敏感性复核。")
    status, interpretation = readiness_status(score)
    return {
        "score": round(score, 2),
        "status": status,
        "interpretation": interpretation,
        "remaining_gaps": gaps[:4],
        "dimension_scores": {key: round(value, 2) for key, value in checks.items()},
    }


def build_strategic_hygiene(request: dict[str, Any], commitments: list[dict[str, Any]], framework_selection: dict[str, Any]) -> dict[str, Any]:
    if request.get("strategic_hygiene"):
        return request["strategic_hygiene"]
    payoff_basis_ok = bool(request.get("payoff_matrix")) and all(row.get("basis") for row in request.get("payoff_matrix", []))
    reaction_rationale_ok = all(likely.get("rationale") for item in request.get("reaction_estimates", []) for likely in item.get("likely_actions", []))
    commitment_ok = bool(commitments) and max(number(item.get("score"), 0.0) for item in commitments) >= 0.75
    checks = [
        {
            "name": "payoff hygiene",
            "status": "pass" if payoff_basis_ok else "watch",
            "note": "payoff 已标记 observed / estimated / assumed。" if payoff_basis_ok else "仍需标记 payoff 来源，避免把估计当事实。",
        },
        {
            "name": "reaction hygiene",
            "status": "pass" if reaction_rationale_ok else "watch",
            "note": "对手反应均有激励理由。" if reaction_rationale_ok else "部分对手反应缺少激励解释。",
        },
        {
            "name": "commitment hygiene",
            "status": "pass" if commitment_ok else "watch",
            "note": "至少一个关键承诺通过可信度门槛。" if commitment_ok else "承诺仍可能被视为便宜话术。",
        },
        {
            "name": "framework hygiene",
            "status": "pass" if framework_selection.get("primary") else "watch",
            "note": "框架选择先从结构出发，而不是只套用著名博弈。" if framework_selection.get("primary") else "需要说明主框架为何适配。",
        },
        {
            "name": "safety hygiene",
            "status": "pass" if request.get("warnings") else "watch",
            "note": "已标记法律、投资、反垄断或合规边界。" if request.get("warnings") else "需要补充合规和高风险边界。",
        },
    ]
    return {"checks": checks}


def build_next_information(request: dict[str, Any], readiness: dict[str, Any], sensitivity: dict[str, Any]) -> dict[str, Any]:
    if request.get("next_information"):
        return request["next_information"]
    priorities = [
        {
            "item": "竞品真实降价承受能力",
            "why_it_matters": "决定竞品跟随降价是否是可持续威胁，而不是短期信号。",
            "collection_method": "观察竞品报价、销售话术、续费折扣和公开成本信号。",
        },
        {
            "item": "核心渠道对年度合作包的接受阈值",
            "why_it_matters": "决定渠道绑定是否能成为可信承诺。",
            "collection_method": "用 2-3 个核心渠道做返点、线索、交付支持的访谈和报价测试。",
        },
        {
            "item": "免费版本是否有真实资源投入",
            "why_it_matters": "决定竞品免费版是便宜信号还是可持续进攻。",
            "collection_method": "追踪产品可用性、渠道推广、转化路径和续费机制。",
        },
    ]
    for gap in readiness.get("remaining_gaps", []):
        priorities.append({"item": "策略准备度缺口", "why_it_matters": gap, "collection_method": "补齐后重新生成博弈报告。"})
    if sensitivity.get("stability") == "fragile":
        priorities.insert(
            0,
            {
                "item": "最危险假设验证",
                "why_it_matters": text(sensitivity.get("most_dangerous_assumption")),
                "collection_method": "先做小规模渠道或客户测试，再决定是否推进不可逆动作。",
            },
        )
    return {"priorities": priorities[:5]}


def build_report(request: dict[str, Any]) -> dict[str, Any]:
    actions = index_by_id(request.get("our_actions", []))
    payoffs = expected_payoffs(request)
    commitments = enrich_commitments(request)
    framework_selection = request.get("framework_selection") or infer_framework_selection(request)
    recommendation = request.get("recommendation", {})
    primary_id = text(recommendation.get("primary_action"), "") or (payoffs[0]["action_id"] if payoffs else "")
    secondary_id = text(recommendation.get("secondary_action"), "")
    avoid_id = text(recommendation.get("avoid_action"), "") or (payoffs[-1]["action_id"] if payoffs else "")
    primary_name = label(primary_id, actions, primary_id)
    secondary_name = label(secondary_id, actions, secondary_id) if secondary_id else ""
    avoid_name = label(avoid_id, actions, avoid_id)
    best_commitment = next((item for item in commitments if text(item.get("action_id"), "") == primary_id), None)
    commitment_verdict = (
        f"{text(best_commitment.get('commitment'))}：{best_commitment['level']}（{fmt_pct(best_commitment['score'])}）"
        if best_commitment
        else "未提供主推荐动作的承诺评分"
    )
    reaction_summary = top_reactions(request, primary_id)
    top_reaction = reaction_summary[0] if reaction_summary else None
    one_sentence = (
        f"建议以「{primary_name}」为主"
        + (f"、以「{secondary_name}」为辅" if secondary_name else "")
        + f"，避免「{avoid_name}」；"
        + (f"主反应风险是「{top_reaction['action']}」（约 {fmt_pct(top_reaction['probability'])}）。" if top_reaction else "当前反应数据不足，需继续补齐对手动作。")
    )
    sensitivity = build_sensitivity(request, payoffs, commitments, primary_id)
    readiness = build_strategy_readiness(request, payoffs, commitments, framework_selection, sensitivity)
    strategic_hygiene = build_strategic_hygiene(request, commitments, framework_selection)
    next_information = build_next_information(request, readiness, sensitivity)
    return {
        "title": text(request.get("title"), "博弈论策略报告"),
        "generated_at": str(date.today()),
        "case_context": request.get("case_context", {}),
        "framework_selection": framework_selection,
        "summary": {
            "one_sentence": one_sentence,
            "recommended_action": primary_name,
            "secondary_action": secondary_name,
            "avoid_action": avoid_name,
            "reason": text(recommendation.get("reason"), "基于当前 payoff、对手反应和承诺可信度的综合判断。"),
            "commitment_verdict": commitment_verdict,
            "strategy_readiness": readiness["score"],
            "stability": sensitivity["stability"],
        },
        "players": request.get("players", []),
        "our_actions": request.get("our_actions", []),
        "opponent_actions": request.get("opponent_actions", []),
        "reaction_estimates": request.get("reaction_estimates", []),
        "top_reactions": reaction_summary,
        "payoff_matrix": request.get("payoff_matrix", []),
        "expected_payoffs": payoffs,
        "commitment_tests": commitments,
        "strategy_readiness": readiness,
        "strategic_hygiene": strategic_hygiene,
        "sensitivity": sensitivity,
        "next_information": next_information,
        "signals": request.get("signals", []),
        "equilibrium": request.get("equilibrium", {}),
        "repeated_game": request.get("repeated_game", {}),
        "rounds": request.get("rounds", []),
        "scenario_triggers": request.get("scenario_triggers", []),
        "warnings": request.get("warnings", []),
    }


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    if not rows:
        return "_暂无数据。_"
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(pipe_escape(cell) for cell in row) + " |")
    return "\n".join(lines)


def render_markdown(report: dict[str, Any]) -> str:
    context = report.get("case_context", {})
    frameworks = report.get("framework_selection", {})
    md: list[str] = [
        f"# {report['title']}",
        "",
        f"> 生成日期：{report['generated_at']}",
        "",
        "## 一句话结论",
        "",
        report["summary"]["one_sentence"],
        "",
        "## 推荐动作",
        "",
        f"- 主动作：{report['summary']['recommended_action']}",
        f"- 辅助动作：{report['summary']['secondary_action'] or '无'}",
        f"- 避免动作：{report['summary']['avoid_action']}",
        f"- 推荐理由：{report['summary']['reason']}",
        f"- 承诺判断：{report['summary']['commitment_verdict']}",
        "",
        "## 问题定义",
        "",
        f"- 决策问题：{text(context.get('decision_question'))}",
        f"- 时间范围：{text(context.get('time_horizon'))}",
        f"- 成功标准：{text(context.get('success_metric'))}",
        "",
        "## 框架选择与组合",
        "",
        f"- 主框架：{text(frameworks.get('primary'))}",
        f"- 组合逻辑：{text(frameworks.get('combination_logic'))}",
        "",
        "### 辅助镜头",
        "",
    ]
    for lens in frameworks.get("secondary_lenses", []):
        md.append(f"- {text(lens)}")
    md.extend(
        [
            "",
            "### 排除或降级的框架",
            "",
            md_table(
                ["框架", "原因"],
                [[item.get("name"), item.get("reason")] for item in frameworks.get("excluded_frameworks", [])],
            ),
            "",
            "## 策略准备度",
            "",
            f"- 准备度：{fmt_pct(report.get('strategy_readiness', {}).get('score'))}",
            f"- 状态：{text(report.get('strategy_readiness', {}).get('status'))}",
            f"- 判断：{text(report.get('strategy_readiness', {}).get('interpretation'))}",
            "",
            "### 剩余缺口",
            "",
        ]
    )
    for gap in report.get("strategy_readiness", {}).get("remaining_gaps", []):
        md.append(f"- {text(gap)}")
    md.extend(
        [
            "",
            "## 战略卫生检查",
            "",
            md_table(
                ["检查", "状态", "说明"],
                [[item.get("name"), item.get("status"), item.get("note")] for item in report.get("strategic_hygiene", {}).get("checks", [])],
            ),
            "",
            "## 敏感性与稳定性",
            "",
            f"- 稳定性：{text(report.get('sensitivity', {}).get('stability'))}",
            f"- 收益差距：{fmt_num(report.get('sensitivity', {}).get('payoff_gap'))}",
            f"- 最危险假设：{text(report.get('sensitivity', {}).get('most_dangerous_assumption'))}",
            "",
            md_table(
                ["压力测试", "假设变化", "影响", "对推荐的影响"],
                [
                    [item.get("name"), item.get("assumption_shift"), item.get("expected_effect"), item.get("recommendation_impact")]
                    for item in report.get("sensitivity", {}).get("stress_tests", [])
                ],
            ),
            "",
            "## 下一步信息",
            "",
            md_table(
                ["信息", "为什么重要", "收集方法"],
                [[item.get("item"), item.get("why_it_matters"), item.get("collection_method")] for item in report.get("next_information", {}).get("priorities", [])],
            ),
            "",
        "## 玩家地图",
        "",
        md_table(
            ["玩家", "角色", "目标", "约束"],
            [[p.get("name"), p.get("role"), text(p.get("objectives")), text(p.get("constraints"))] for p in report.get("players", [])],
        ),
        "",
        "## 策略集合",
        "",
        md_table(
            ["我们的动作", "说明", "数据基础"],
            [[a.get("name"), a.get("description"), a.get("data_basis")] for a in report.get("our_actions", [])],
        ),
        "",
        "## 对手反应地图",
        "",
        ]
    )
    reaction_rows = []
    action_map = index_by_id(report.get("our_actions", []))
    opponent_action_map = index_by_id(report.get("opponent_actions", []))
    for estimate in report.get("reaction_estimates", []):
        for likely in estimate.get("likely_actions", []):
            reaction_rows.append(
                [
                    text(estimate.get("opponent")),
                    label(estimate.get("if_we"), action_map),
                    label(likely.get("action"), opponent_action_map),
                    fmt_pct(likely.get("probability")),
                    likely.get("rationale"),
                ]
            )
    md.append(md_table(["对手", "如果我们", "可能动作", "概率", "理由"], reaction_rows))
    md.extend(
        [
            "",
            "## Payoff 矩阵",
            "",
            md_table(
                ["我们的动作", "对方动作", "我方收益", "对方收益", "渠道收益", "基础", "说明"],
                [
                    [
                        label(row.get("our_action"), action_map),
                        label(row.get("opponent_action"), opponent_action_map, text(row.get("opponent_action"))),
                        fmt_num(row.get("our_payoff")),
                        fmt_num(row.get("opponent_payoff")),
                        fmt_num(row.get("channel_payoff")),
                        row.get("basis"),
                        row.get("notes"),
                    ]
                    for row in report.get("payoff_matrix", [])
                ],
            ),
            "",
            "## 预期收益排序",
            "",
            md_table(
                ["动作", "预期收益", "计算方式"],
                [[item["action_name"], fmt_num(item["expected_payoff"]), item["method"]] for item in report.get("expected_payoffs", [])],
            ),
            "",
            "## 承诺与信号可信度",
            "",
            md_table(
                ["动作", "承诺", "可信度", "评分", "证据基础", "说明"],
                [
                    [item.get("action_name"), item.get("commitment"), item.get("level"), fmt_pct(item.get("score")), item.get("evidence_basis"), item.get("notes")]
                    for item in report.get("commitment_tests", [])
                ],
            ),
            "",
        ]
    )
    if report.get("signals"):
        md.extend(
            [
                "### 外部信号",
                "",
                md_table(
                    ["发送方", "信号", "质量", "解释"],
                    [[s.get("sender"), s.get("signal"), s.get("quality"), s.get("interpretation")] for s in report.get("signals", [])],
                ),
                "",
            ]
        )
    equilibrium = report.get("equilibrium", {})
    md.extend(
        [
            "## 均衡解释",
            "",
            f"- 主模型：{text(equilibrium.get('primary_frame'))}",
            f"- 解释：{text(equilibrium.get('interpretation'))}",
            "",
        ]
    )
    for item in equilibrium.get("candidate_equilibria", []):
        md.append(f"- {text(item)}")
    repeated = report.get("repeated_game", {})
    md.extend(
        [
            "",
            "## 重复博弈与关系动态",
            "",
            text(repeated.get("interpretation")),
            "",
        ]
    )
    for risk in repeated.get("relationship_risks", []):
        md.append(f"- {text(risk)}")
    if report.get("rounds"):
        md.extend(
            [
                "",
                "## 动态更新日志",
                "",
                md_table(
                    ["轮次", "日期", "阶段", "新增信息", "判断", "状态"],
                    [
                        [r.get("round"), r.get("date"), r.get("stage"), text(r.get("new_information")), r.get("interim_judgment"), r.get("recommendation_status")]
                        for r in report.get("rounds", [])
                    ],
                ),
                "",
            ]
        )
    md.extend(
        [
            "## 需要重新打开报告的触发器",
            "",
            md_table(
                ["触发器", "更新规则"],
                [[item.get("trigger"), item.get("update_rule")] for item in report.get("scenario_triggers", [])],
            ),
            "",
            "## 风险与边界",
            "",
        ]
    )
    for warning in report.get("warnings", []):
        md.append(f"- {text(warning)}")
    md.extend(
        [
            "",
            "## 自动化说明",
            "",
            "本报告由 `yao-gametheory-skill` 从同一份结构化 JSON 自动生成，Markdown、HTML、DOCX、PDF 与 canonical JSON 应保持同步。",
            "",
        ]
    )
    return "\n".join(md)


CSS = """
:root {
  --ink: #17211b;
  --muted: #5c675f;
  --line: #d9e0d8;
  --paper: #fbfcf8;
  --surface: #ffffff;
  --accent: #196f5a;
  --accent-soft: #e6f3ef;
  --risk: #9f3a2f;
  --warn: #8a6718;
}
* { box-sizing: border-box; }
body {
  margin: 0;
  color: var(--ink);
  background: var(--paper);
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", sans-serif;
  line-height: 1.65;
}
.topbar {
  position: sticky;
  top: 0;
  z-index: 10;
  display: flex;
  justify-content: space-between;
  gap: 16px;
  align-items: center;
  padding: 10px 24px;
  border-bottom: 1px solid var(--line);
  background: rgba(251, 252, 248, 0.95);
}
.nav { display: flex; flex-wrap: wrap; gap: 10px; font-size: 13px; }
.nav a { color: var(--muted); text-decoration: none; }
.actions button {
  border: 1px solid var(--line);
  background: var(--surface);
  color: var(--ink);
  border-radius: 6px;
  padding: 7px 10px;
  cursor: pointer;
}
main { max-width: 1120px; margin: 0 auto; padding: 34px 24px 72px; }
.hero {
  padding: 24px 0 18px;
  border-bottom: 2px solid var(--ink);
}
.eyebrow { color: var(--accent); font-weight: 700; font-size: 13px; letter-spacing: 0; }
h1 { margin: 8px 0 10px; font-size: 34px; line-height: 1.18; letter-spacing: 0; }
h2 { margin: 34px 0 12px; font-size: 22px; letter-spacing: 0; }
h3 { margin: 22px 0 10px; font-size: 17px; letter-spacing: 0; }
.summary {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
  gap: 12px;
  margin: 18px 0 6px;
}
.metric {
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--surface);
  padding: 14px;
}
.metric .label { color: var(--muted); font-size: 12px; }
.metric .value { font-size: 18px; font-weight: 750; margin-top: 4px; }
.callout {
  margin: 18px 0;
  border-left: 4px solid var(--accent);
  background: var(--accent-soft);
  padding: 14px 16px;
}
.warn { border-left-color: var(--warn); background: #fff8e5; }
.risk { border-left-color: var(--risk); background: #fff0ed; }
.table-wrap { overflow-x: auto; border: 1px solid var(--line); border-radius: 8px; background: var(--surface); }
table { width: 100%; border-collapse: collapse; min-width: 760px; }
th, td {
  padding: 10px 12px;
  border-bottom: 1px solid var(--line);
  text-align: left;
  vertical-align: top;
  overflow-wrap: anywhere;
  word-break: break-word;
}
th { background: #edf3ec; font-size: 13px; }
tr:last-child td { border-bottom: 0; }
.pill {
  display: inline-block;
  border-radius: 999px;
  padding: 2px 8px;
  background: #edf3ec;
  color: var(--accent);
  font-size: 12px;
  font-weight: 700;
}
details {
  border: 1px solid var(--line);
  border-radius: 8px;
  background: var(--surface);
  padding: 12px 14px;
  margin-top: 12px;
}
summary { cursor: pointer; font-weight: 750; }
footer { color: var(--muted); font-size: 13px; margin-top: 42px; border-top: 1px solid var(--line); padding-top: 18px; }
@media (max-width: 760px) {
  .topbar { position: static; align-items: flex-start; flex-direction: column; }
  main { padding: 24px 16px 56px; }
  h1 { font-size: 27px; }
  .summary { grid-template-columns: 1fr; }
}
@media print {
  @page { size: A4 landscape; margin: 12mm; }
  .topbar, .actions { display: none !important; }
  body { background: #fff; font-size: 10pt; line-height: 1.45; }
  main { max-width: none; padding: 0; width: 100%; }
  h1 { font-size: 22pt; }
  h2 { font-size: 15pt; margin-top: 18px; }
  h3 { font-size: 12pt; margin-top: 14px; }
  .summary { grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 8px; }
  .metric { padding: 8px; }
  .metric .value { font-size: 11pt; overflow-wrap: anywhere; }
  .callout { padding: 8px 10px; margin: 10px 0; }
  .table-wrap { overflow: visible; border-radius: 4px; break-inside: avoid; page-break-inside: avoid; }
  table { min-width: 0; width: 100%; table-layout: fixed; font-size: 8.5pt; }
  th, td { padding: 5px 6px; white-space: normal; overflow-wrap: anywhere; word-break: break-word; }
  th { font-size: 8.5pt; }
  tr { break-inside: avoid; page-break-inside: avoid; }
  details { break-inside: avoid; }
}
"""


def html_table(headers: list[str], rows: list[list[Any]]) -> str:
    if not rows:
        return "<p>暂无数据。</p>"
    header_html = "".join(f"<th>{escape(head)}</th>" for head in headers)
    body_rows = []
    for row in rows:
        cells = "".join(f"<td>{h(cell)}</td>" for cell in row)
        body_rows.append(f"<tr>{cells}</tr>")
    return f'<div class="table-wrap"><table><thead><tr>{header_html}</tr></thead><tbody>{"".join(body_rows)}</tbody></table></div>'


def render_html(report: dict[str, Any]) -> str:
    context = report.get("case_context", {})
    frameworks = report.get("framework_selection", {})
    readiness = report.get("strategy_readiness", {})
    hygiene = report.get("strategic_hygiene", {})
    sensitivity = report.get("sensitivity", {})
    next_information = report.get("next_information", {})
    action_map = index_by_id(report.get("our_actions", []))
    opponent_action_map = index_by_id(report.get("opponent_actions", []))
    reaction_rows = []
    for estimate in report.get("reaction_estimates", []):
        for likely in estimate.get("likely_actions", []):
            reaction_rows.append(
                [
                    estimate.get("opponent"),
                    label(estimate.get("if_we"), action_map),
                    label(likely.get("action"), opponent_action_map),
                    fmt_pct(likely.get("probability")),
                    likely.get("rationale"),
                ]
            )
    payoff_rows = [
        [
            label(row.get("our_action"), action_map),
            label(row.get("opponent_action"), opponent_action_map, text(row.get("opponent_action"))),
            fmt_num(row.get("our_payoff")),
            fmt_num(row.get("opponent_payoff")),
            fmt_num(row.get("channel_payoff")),
            row.get("basis"),
            row.get("notes"),
        ]
        for row in report.get("payoff_matrix", [])
    ]
    warning_items = "".join(f"<li>{h(item)}</li>" for item in report.get("warnings", []))
    equilibrium = report.get("equilibrium", {})
    candidate_items = "".join(f"<li>{h(item)}</li>" for item in equilibrium.get("candidate_equilibria", []))
    repeated = report.get("repeated_game", {})
    repeated_items = "".join(f"<li>{h(item)}</li>" for item in repeated.get("relationship_risks", []))
    signal_rows = [[s.get("sender"), s.get("signal"), s.get("quality"), s.get("interpretation")] for s in report.get("signals", [])]
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{h(report['title'])}</title>
  <style>{CSS}</style>
</head>
<body>
  <div class="topbar">
    <nav class="nav">
      <a href="#summary">结论</a>
      <a href="#frameworks">框架</a>
      <a href="#readiness">准备度</a>
      <a href="#sensitivity">敏感性</a>
      <a href="#players">玩家</a>
      <a href="#reactions">反应</a>
      <a href="#payoff">Payoff</a>
      <a href="#commitment">承诺</a>
      <a href="#equilibrium">均衡</a>
      <a href="#updates">更新</a>
      <a href="#risks">风险</a>
    </nav>
    <div class="actions">
      <button onclick="expandAndPrint()">Print / Save as PDF</button>
    </div>
  </div>
  <main>
    <section class="hero" id="summary">
      <div class="eyebrow">Yao Game Theory Skill</div>
      <h1>{h(report['title'])}</h1>
      <p>{h(report['summary']['one_sentence'])}</p>
      <div class="summary">
        <div class="metric"><div class="label">主动作</div><div class="value">{h(report['summary']['recommended_action'])}</div></div>
        <div class="metric"><div class="label">辅助动作</div><div class="value">{h(report['summary']['secondary_action'] or '无')}</div></div>
        <div class="metric"><div class="label">避免动作</div><div class="value">{h(report['summary']['avoid_action'])}</div></div>
        <div class="metric"><div class="label">承诺判断</div><div class="value">{h(report['summary']['commitment_verdict'])}</div></div>
        <div class="metric"><div class="label">策略准备度</div><div class="value">{fmt_pct(readiness.get('score'))}</div></div>
        <div class="metric"><div class="label">稳定性</div><div class="value">{h(sensitivity.get('stability'))}</div></div>
      </div>
    </section>

    <section>
      <h2>推荐动作</h2>
      <div class="callout">{h(report['summary']['reason'])}</div>
      <p><span class="pill">决策问题</span> {h(context.get('decision_question'))}</p>
      <p><span class="pill">时间范围</span> {h(context.get('time_horizon'))}</p>
      <p><span class="pill">成功标准</span> {h(context.get('success_metric'))}</p>
    </section>

    <section id="frameworks">
      <h2>框架选择与组合</h2>
      <p><span class="pill">主框架</span> {h(frameworks.get('primary'))}</p>
      <div class="callout">{h(frameworks.get('combination_logic'))}</div>
      <h3>辅助镜头</h3>
      <ul>{"".join(f"<li>{h(item)}</li>" for item in frameworks.get("secondary_lenses", []))}</ul>
      <h3>排除或降级的框架</h3>
      {html_table(["框架", "原因"], [[item.get("name"), item.get("reason")] for item in frameworks.get("excluded_frameworks", [])])}
    </section>

    <section id="readiness">
      <h2>策略准备度</h2>
      <p><span class="pill">准备度</span> {fmt_pct(readiness.get('score'))} <span class="pill">状态</span> {h(readiness.get('status'))}</p>
      <div class="callout">{h(readiness.get('interpretation'))}</div>
      <h3>剩余缺口</h3>
      <ul>{"".join(f"<li>{h(item)}</li>" for item in readiness.get("remaining_gaps", []))}</ul>
      <h3>战略卫生检查</h3>
      {html_table(["检查", "状态", "说明"], [[item.get("name"), item.get("status"), item.get("note")] for item in hygiene.get("checks", [])])}
    </section>

    <section id="sensitivity">
      <h2>敏感性与稳定性</h2>
      <p><span class="pill">稳定性</span> {h(sensitivity.get('stability'))} <span class="pill">收益差距</span> {fmt_num(sensitivity.get('payoff_gap'))}</p>
      <div class="callout warn">{h(sensitivity.get('most_dangerous_assumption'))}</div>
      {html_table(["压力测试", "假设变化", "影响", "对推荐的影响"], [[item.get("name"), item.get("assumption_shift"), item.get("expected_effect"), item.get("recommendation_impact")] for item in sensitivity.get("stress_tests", [])])}
      <h3>下一步信息</h3>
      {html_table(["信息", "为什么重要", "收集方法"], [[item.get("item"), item.get("why_it_matters"), item.get("collection_method")] for item in next_information.get("priorities", [])])}
    </section>

    <section id="players">
      <h2>玩家地图</h2>
      {html_table(["玩家", "角色", "目标", "约束"], [[p.get("name"), p.get("role"), text(p.get("objectives")), text(p.get("constraints"))] for p in report.get("players", [])])}
      <h3>我们的策略集合</h3>
      {html_table(["动作", "说明", "数据基础"], [[a.get("name"), a.get("description"), a.get("data_basis")] for a in report.get("our_actions", [])])}
    </section>

    <section id="reactions">
      <h2>对手反应地图</h2>
      {html_table(["对手", "如果我们", "可能动作", "概率", "理由"], reaction_rows)}
    </section>

    <section id="payoff">
      <h2>Payoff 矩阵与排序</h2>
      {html_table(["我们的动作", "对方动作", "我方收益", "对方收益", "渠道收益", "基础", "说明"], payoff_rows)}
      <h3>预期收益排序</h3>
      {html_table(["动作", "预期收益", "计算方式"], [[item["action_name"], fmt_num(item["expected_payoff"]), item["method"]] for item in report.get("expected_payoffs", [])])}
    </section>

    <section id="commitment">
      <h2>承诺与信号可信度</h2>
      {html_table(["动作", "承诺", "可信度", "评分", "证据基础", "说明"], [[item.get("action_name"), item.get("commitment"), item.get("level"), fmt_pct(item.get("score")), item.get("evidence_basis"), item.get("notes")] for item in report.get("commitment_tests", [])])}
      <h3>外部信号</h3>
      {html_table(["发送方", "信号", "质量", "解释"], signal_rows)}
    </section>

    <section id="equilibrium">
      <h2>均衡解释</h2>
      <p><span class="pill">主模型</span> {h(equilibrium.get('primary_frame'))}</p>
      <p>{h(equilibrium.get('interpretation'))}</p>
      <ul>{candidate_items}</ul>
      <h3>重复博弈与关系动态</h3>
      <p>{h(repeated.get('interpretation'))}</p>
      <ul>{repeated_items}</ul>
    </section>

    <section id="updates">
      <h2>动态更新日志</h2>
      {html_table(["轮次", "日期", "阶段", "新增信息", "判断", "状态"], [[r.get("round"), r.get("date"), r.get("stage"), text(r.get("new_information")), r.get("interim_judgment"), r.get("recommendation_status")] for r in report.get("rounds", [])])}
      <h3>重新打开报告的触发器</h3>
      {html_table(["触发器", "更新规则"], [[item.get("trigger"), item.get("update_rule")] for item in report.get("scenario_triggers", [])])}
    </section>

    <section id="risks">
      <h2>风险与边界</h2>
      <div class="callout risk"><ul>{warning_items}</ul></div>
      <details>
        <summary>附录：自动化说明</summary>
        <p>本报告由同一份结构化 JSON 自动生成，Markdown、HTML、DOCX、PDF 与 canonical JSON 应保持同步。后续对手动作可通过 update JSON 合并后重新导出。</p>
      </details>
    </section>
    <footer>Generated by yao-gametheory-skill on {h(report['generated_at'])}.</footer>
  </main>
  <script>
    function expandAndPrint() {{
      document.querySelectorAll('details').forEach(function(item) {{ item.open = true; }});
      window.print();
    }}
  </script>
</body>
</html>
"""
    return html


def docx_lines(report: dict[str, Any]) -> list[tuple[str, str]]:
    frameworks = report.get("framework_selection", {})
    lines: list[tuple[str, str]] = [
        ("title", report["title"]),
        ("normal", f"生成日期：{report['generated_at']}"),
        ("heading", "一句话结论"),
        ("normal", report["summary"]["one_sentence"]),
        ("heading", "推荐动作"),
        ("bullet", f"主动作：{report['summary']['recommended_action']}"),
        ("bullet", f"辅助动作：{report['summary']['secondary_action'] or '无'}"),
        ("bullet", f"避免动作：{report['summary']['avoid_action']}"),
        ("bullet", f"推荐理由：{report['summary']['reason']}"),
        ("bullet", f"承诺判断：{report['summary']['commitment_verdict']}"),
        ("heading", "框架选择与组合"),
        ("normal", f"主框架：{text(frameworks.get('primary'))}"),
        ("normal", f"组合逻辑：{text(frameworks.get('combination_logic'))}"),
        ("heading", "辅助镜头"),
    ]
    for lens in frameworks.get("secondary_lenses", []):
        lines.append(("bullet", text(lens)))
    lines.extend(
        [
            ("heading", "排除或降级的框架"),
        ]
    )
    for item in frameworks.get("excluded_frameworks", []):
        lines.append(("bullet", f"{text(item.get('name'))}：{text(item.get('reason'))}"))
    lines.extend(
        [
        ("heading", "玩家地图"),
        ]
    )
    for p in report.get("players", []):
        lines.append(("normal", f"{text(p.get('name'))}｜{text(p.get('role'))}｜目标：{text(p.get('objectives'))}｜约束：{text(p.get('constraints'))}"))
    lines.append(("heading", "对手反应地图"))
    action_map = index_by_id(report.get("our_actions", []))
    opponent_action_map = index_by_id(report.get("opponent_actions", []))
    for estimate in report.get("reaction_estimates", []):
        for likely in estimate.get("likely_actions", []):
            lines.append(
                (
                    "normal",
                    f"{text(estimate.get('opponent'))}｜如果我们：{label(estimate.get('if_we'), action_map)}｜可能动作：{label(likely.get('action'), opponent_action_map)}｜概率：{fmt_pct(likely.get('probability'))}｜{text(likely.get('rationale'))}",
                )
            )
    lines.append(("heading", "Payoff 矩阵"))
    for row in report.get("payoff_matrix", []):
        lines.append(
            (
                "normal",
                f"{label(row.get('our_action'), action_map)} x {label(row.get('opponent_action'), opponent_action_map, text(row.get('opponent_action')))}｜我方 {fmt_num(row.get('our_payoff'))}｜对方 {fmt_num(row.get('opponent_payoff'))}｜渠道 {fmt_num(row.get('channel_payoff'))}｜{text(row.get('notes'))}",
            )
        )
    lines.append(("heading", "承诺可信度"))
    for item in report.get("commitment_tests", []):
        lines.append(("normal", f"{text(item.get('action_name'))}｜{text(item.get('commitment'))}｜{text(item.get('level'))}｜{fmt_pct(item.get('score'))}｜{text(item.get('notes'))}"))
    lines.append(("heading", "动态更新日志"))
    for r in report.get("rounds", []):
        lines.append(("normal", f"第 {text(r.get('round'))} 轮｜{text(r.get('date'))}｜{text(r.get('stage'))}｜{text(r.get('interim_judgment'))}｜{text(r.get('recommendation_status'))}"))
    lines.append(("heading", "风险与边界"))
    for warning in report.get("warnings", []):
        lines.append(("bullet", text(warning)))
    return lines


def simple_docx_paragraph(style: str, value: str) -> str:
    safe = escape(value)
    if style == "title":
        props = "<w:pPr><w:jc w:val=\"center\"/><w:spacing w:after=\"220\"/></w:pPr>"
        run_props = "<w:rPr><w:rFonts w:ascii=\"Arial\" w:hAnsi=\"Arial\" w:eastAsia=\"Microsoft YaHei\"/><w:b/><w:sz w:val=\"34\"/></w:rPr>"
    elif style == "heading":
        props = "<w:pPr><w:spacing w:before=\"220\" w:after=\"100\"/></w:pPr>"
        run_props = "<w:rPr><w:rFonts w:ascii=\"Arial\" w:hAnsi=\"Arial\" w:eastAsia=\"Microsoft YaHei\"/><w:b/><w:sz w:val=\"26\"/></w:rPr>"
    elif style == "bullet":
        props = "<w:pPr><w:spacing w:after=\"40\"/></w:pPr>"
        run_props = "<w:rPr><w:rFonts w:ascii=\"Arial\" w:hAnsi=\"Arial\" w:eastAsia=\"Microsoft YaHei\"/><w:sz w:val=\"20\"/></w:rPr>"
        safe = "• " + safe
    else:
        props = "<w:pPr><w:spacing w:after=\"70\"/></w:pPr>"
        run_props = "<w:rPr><w:rFonts w:ascii=\"Arial\" w:hAnsi=\"Arial\" w:eastAsia=\"Microsoft YaHei\"/><w:sz w:val=\"20\"/></w:rPr>"
    return f"<w:p>{props}<w:r>{run_props}<w:t xml:space=\"preserve\">{safe}</w:t></w:r></w:p>"


def simple_docx_cell(value: Any, width: int, bold: bool = False) -> str:
    run_props = "<w:rPr><w:rFonts w:ascii=\"Arial\" w:hAnsi=\"Arial\" w:eastAsia=\"Microsoft YaHei\"/><w:sz w:val=\"18\"/>"
    if bold:
        run_props += "<w:b/>"
    run_props += "</w:rPr>"
    safe = escape(text(value))
    return (
        "<w:tc>"
        f"<w:tcPr><w:tcW w:w=\"{width}\" w:type=\"dxa\"/><w:vAlign w:val=\"top\"/></w:tcPr>"
        "<w:p><w:pPr><w:spacing w:after=\"0\"/></w:pPr>"
        f"<w:r>{run_props}<w:t xml:space=\"preserve\">{safe}</w:t></w:r>"
        "</w:p></w:tc>"
    )


def simple_docx_table(headers: list[str], rows: list[list[Any]], widths: list[int] | None = None) -> str:
    if not rows:
        return simple_docx_paragraph("normal", "暂无数据。")
    total_width = 15400
    widths = widths or [total_width // len(headers)] * len(headers)
    if len(widths) != len(headers):
        widths = [total_width // len(headers)] * len(headers)
    grid = "".join(f"<w:gridCol w:w=\"{width}\"/>" for width in widths)
    header_cells = "".join(simple_docx_cell(head, widths[idx], bold=True) for idx, head in enumerate(headers))
    body_rows = []
    for row in rows:
        padded = list(row)[: len(headers)] + [""] * max(0, len(headers) - len(row))
        cells = "".join(simple_docx_cell(padded[idx], widths[idx]) for idx in range(len(headers)))
        body_rows.append(f"<w:tr>{cells}</w:tr>")
    return (
        "<w:tbl>"
        "<w:tblPr>"
        "<w:tblW w:w=\"15400\" w:type=\"dxa\"/>"
        "<w:tblLayout w:type=\"fixed\"/>"
        "<w:tblBorders>"
        "<w:top w:val=\"single\" w:sz=\"4\" w:color=\"D9E0D8\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:color=\"D9E0D8\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:color=\"D9E0D8\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:color=\"D9E0D8\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:color=\"D9E0D8\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:color=\"D9E0D8\"/>"
        "</w:tblBorders>"
        "<w:tblCellMar>"
        "<w:top w:w=\"80\" w:type=\"dxa\"/><w:left w:w=\"80\" w:type=\"dxa\"/>"
        "<w:bottom w:w=\"80\" w:type=\"dxa\"/><w:right w:w=\"80\" w:type=\"dxa\"/>"
        "</w:tblCellMar>"
        "</w:tblPr>"
        f"<w:tblGrid>{grid}</w:tblGrid>"
        f"<w:tr>{header_cells}</w:tr>"
        f"{''.join(body_rows)}"
        "</w:tbl>"
    )


def simple_docx_blocks(report: dict[str, Any]) -> list[str]:
    frameworks = report.get("framework_selection", {})
    readiness = report.get("strategy_readiness", {})
    hygiene = report.get("strategic_hygiene", {})
    sensitivity = report.get("sensitivity", {})
    next_information = report.get("next_information", {})
    action_map = index_by_id(report.get("our_actions", []))
    opponent_action_map = index_by_id(report.get("opponent_actions", []))
    blocks = [
        simple_docx_paragraph("title", report["title"]),
        simple_docx_paragraph("normal", f"生成日期：{report['generated_at']}"),
        simple_docx_paragraph("heading", "一句话结论"),
        simple_docx_paragraph("normal", report["summary"]["one_sentence"]),
        simple_docx_paragraph("heading", "推荐动作"),
        simple_docx_paragraph("bullet", f"主动作：{report['summary']['recommended_action']}"),
        simple_docx_paragraph("bullet", f"辅助动作：{report['summary']['secondary_action'] or '无'}"),
        simple_docx_paragraph("bullet", f"避免动作：{report['summary']['avoid_action']}"),
        simple_docx_paragraph("bullet", f"推荐理由：{report['summary']['reason']}"),
        simple_docx_paragraph("bullet", f"承诺判断：{report['summary']['commitment_verdict']}"),
        simple_docx_paragraph("heading", "框架选择与组合"),
        simple_docx_paragraph("normal", f"主框架：{text(frameworks.get('primary'))}"),
        simple_docx_paragraph("normal", f"组合逻辑：{text(frameworks.get('combination_logic'))}"),
        simple_docx_paragraph("heading", "辅助镜头"),
    ]
    for lens in frameworks.get("secondary_lenses", []):
        blocks.append(simple_docx_paragraph("bullet", text(lens)))
    blocks.extend(
        [
            simple_docx_paragraph("heading", "排除或降级的框架"),
            simple_docx_table(
                ["框架", "原因"],
                [[item.get("name"), item.get("reason")] for item in frameworks.get("excluded_frameworks", [])],
                [3000, 12400],
            ),
            simple_docx_paragraph("heading", "策略准备度"),
            simple_docx_paragraph("normal", f"准备度：{fmt_pct(readiness.get('score'))}｜状态：{text(readiness.get('status'))}"),
            simple_docx_paragraph("normal", f"判断：{text(readiness.get('interpretation'))}"),
            simple_docx_paragraph("heading", "剩余缺口"),
        ]
    )
    for gap in readiness.get("remaining_gaps", []):
        blocks.append(simple_docx_paragraph("bullet", text(gap)))
    blocks.extend(
        [
            simple_docx_paragraph("heading", "战略卫生检查"),
            simple_docx_table(
                ["检查", "状态", "说明"],
                [[item.get("name"), item.get("status"), item.get("note")] for item in hygiene.get("checks", [])],
                [3400, 2200, 9800],
            ),
            simple_docx_paragraph("heading", "敏感性与稳定性"),
            simple_docx_paragraph("normal", f"稳定性：{text(sensitivity.get('stability'))}｜收益差距：{fmt_num(sensitivity.get('payoff_gap'))}"),
            simple_docx_paragraph("normal", f"最危险假设：{text(sensitivity.get('most_dangerous_assumption'))}"),
            simple_docx_table(
                ["压力测试", "假设变化", "影响", "对推荐的影响"],
                [
                    [item.get("name"), item.get("assumption_shift"), item.get("expected_effect"), item.get("recommendation_impact")]
                    for item in sensitivity.get("stress_tests", [])
                ],
                [2500, 4500, 4200, 4200],
            ),
            simple_docx_paragraph("heading", "下一步信息"),
            simple_docx_table(
                ["信息", "为什么重要", "收集方法"],
                [[item.get("item"), item.get("why_it_matters"), item.get("collection_method")] for item in next_information.get("priorities", [])],
                [3600, 5600, 6200],
            ),
            simple_docx_paragraph("heading", "玩家地图"),
            simple_docx_table(
                ["玩家", "角色", "目标", "约束"],
                [[p.get("name"), p.get("role"), text(p.get("objectives")), text(p.get("constraints"))] for p in report.get("players", [])],
                [1800, 2800, 5400, 5400],
            ),
            simple_docx_paragraph("heading", "我们的策略集合"),
            simple_docx_table(
                ["动作", "说明", "数据基础"],
                [[a.get("name"), a.get("description"), a.get("data_basis")] for a in report.get("our_actions", [])],
                [2200, 10600, 2600],
            ),
        ]
    )
    reaction_rows = []
    for estimate in report.get("reaction_estimates", []):
        for likely in estimate.get("likely_actions", []):
            reaction_rows.append(
                [
                    estimate.get("opponent"),
                    label(estimate.get("if_we"), action_map),
                    label(likely.get("action"), opponent_action_map),
                    fmt_pct(likely.get("probability")),
                    likely.get("rationale"),
                ]
            )
    blocks.extend(
        [
            simple_docx_paragraph("heading", "对手反应地图"),
            simple_docx_table(["对手", "如果我们", "可能动作", "概率", "理由"], reaction_rows, [3000, 2200, 2600, 1400, 6200]),
            simple_docx_paragraph("heading", "Payoff 矩阵"),
            simple_docx_table(
                ["我们的动作", "对方动作", "我方", "对方", "渠道", "基础", "说明"],
                [
                    [
                        label(row.get("our_action"), action_map),
                        label(row.get("opponent_action"), opponent_action_map, text(row.get("opponent_action"))),
                        fmt_num(row.get("our_payoff")),
                        fmt_num(row.get("opponent_payoff")),
                        fmt_num(row.get("channel_payoff")),
                        row.get("basis"),
                        row.get("notes"),
                    ]
                    for row in report.get("payoff_matrix", [])
                ],
                [2200, 2600, 1000, 1000, 1000, 1500, 6100],
            ),
            simple_docx_paragraph("heading", "预期收益排序"),
            simple_docx_table(
                ["动作", "预期收益", "计算方式"],
                [[item["action_name"], fmt_num(item["expected_payoff"]), item["method"]] for item in report.get("expected_payoffs", [])],
                [4200, 2800, 8400],
            ),
            simple_docx_paragraph("heading", "承诺可信度"),
            simple_docx_table(
                ["动作", "承诺", "可信度", "评分", "证据基础", "说明"],
                [
                    [item.get("action_name"), item.get("commitment"), item.get("level"), fmt_pct(item.get("score")), item.get("evidence_basis"), item.get("notes")]
                    for item in report.get("commitment_tests", [])
                ],
                [1900, 3300, 1700, 1300, 1800, 5400],
            ),
            simple_docx_paragraph("heading", "动态更新日志"),
            simple_docx_table(
                ["轮次", "日期", "阶段", "新增信息", "判断", "状态"],
                [[r.get("round"), r.get("date"), r.get("stage"), text(r.get("new_information")), r.get("interim_judgment"), r.get("recommendation_status")] for r in report.get("rounds", [])],
                [900, 1800, 1900, 5200, 4400, 1200],
            ),
            simple_docx_paragraph("heading", "风险与边界"),
        ]
    )
    for warning in report.get("warnings", []):
        blocks.append(simple_docx_paragraph("bullet", text(warning)))
    return blocks


def write_simple_docx(path: Path, report: dict[str, Any]) -> None:
    blocks = "\n".join(simple_docx_blocks(report))
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {blocks}
    <w:sectPr>
      <w:pgSz w:w="16838" w:h="11906" w:orient="landscape"/>
      <w:pgMar w:top="720" w:right="720" w:bottom="720" w:left="720" w:header="360" w:footer="360" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>
"""
    rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types_xml)
        docx.writestr("_rels/.rels", rels_xml)
        docx.writestr("word/document.xml", document_xml)


def write_docx(path: Path, report: dict[str, Any]) -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - fallback depends on environment.
        write_simple_docx(path, report)
        return

    document = Document()
    document.add_heading(report["title"], 0)
    document.add_paragraph(f"生成日期：{report['generated_at']}")
    document.add_heading("一句话结论", level=1)
    document.add_paragraph(report["summary"]["one_sentence"])
    document.add_heading("推荐动作", level=1)
    for line in [
        f"主动作：{report['summary']['recommended_action']}",
        f"辅助动作：{report['summary']['secondary_action'] or '无'}",
        f"避免动作：{report['summary']['avoid_action']}",
        f"推荐理由：{report['summary']['reason']}",
        f"承诺判断：{report['summary']['commitment_verdict']}",
    ]:
        document.add_paragraph(line, style="List Bullet")

    def add_table(title: str, headers: list[str], rows: list[list[Any]]) -> None:
        document.add_heading(title, level=1)
        if not rows:
            document.add_paragraph("暂无数据。")
            return
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, head in enumerate(headers):
            table.rows[0].cells[idx].text = head
        for row in rows:
            cells = table.add_row().cells
            for idx, cell in enumerate(row):
                cells[idx].text = text(cell)

    action_map = index_by_id(report.get("our_actions", []))
    opponent_action_map = index_by_id(report.get("opponent_actions", []))
    frameworks = report.get("framework_selection", {})
    readiness = report.get("strategy_readiness", {})
    hygiene = report.get("strategic_hygiene", {})
    sensitivity = report.get("sensitivity", {})
    next_information = report.get("next_information", {})
    document.add_heading("框架选择与组合", level=1)
    document.add_paragraph(f"主框架：{text(frameworks.get('primary'))}")
    document.add_paragraph(f"组合逻辑：{text(frameworks.get('combination_logic'))}")
    for lens in frameworks.get("secondary_lenses", []):
        document.add_paragraph(text(lens), style="List Bullet")
    add_table("排除或降级的框架", ["框架", "原因"], [[item.get("name"), item.get("reason")] for item in frameworks.get("excluded_frameworks", [])])
    document.add_heading("策略准备度", level=1)
    document.add_paragraph(f"准备度：{fmt_pct(readiness.get('score'))}｜状态：{text(readiness.get('status'))}")
    document.add_paragraph(f"判断：{text(readiness.get('interpretation'))}")
    for gap in readiness.get("remaining_gaps", []):
        document.add_paragraph(text(gap), style="List Bullet")
    add_table("战略卫生检查", ["检查", "状态", "说明"], [[item.get("name"), item.get("status"), item.get("note")] for item in hygiene.get("checks", [])])
    document.add_heading("敏感性与稳定性", level=1)
    document.add_paragraph(f"稳定性：{text(sensitivity.get('stability'))}｜收益差距：{fmt_num(sensitivity.get('payoff_gap'))}")
    document.add_paragraph(f"最危险假设：{text(sensitivity.get('most_dangerous_assumption'))}")
    add_table(
        "压力测试",
        ["压力测试", "假设变化", "影响", "对推荐的影响"],
        [[item.get("name"), item.get("assumption_shift"), item.get("expected_effect"), item.get("recommendation_impact")] for item in sensitivity.get("stress_tests", [])],
    )
    add_table(
        "下一步信息",
        ["信息", "为什么重要", "收集方法"],
        [[item.get("item"), item.get("why_it_matters"), item.get("collection_method")] for item in next_information.get("priorities", [])],
    )
    add_table(
        "玩家地图",
        ["玩家", "角色", "目标", "约束"],
        [[p.get("name"), p.get("role"), text(p.get("objectives")), text(p.get("constraints"))] for p in report.get("players", [])],
    )
    reaction_rows = []
    for estimate in report.get("reaction_estimates", []):
        for likely in estimate.get("likely_actions", []):
            reaction_rows.append(
                [estimate.get("opponent"), label(estimate.get("if_we"), action_map), label(likely.get("action"), opponent_action_map), fmt_pct(likely.get("probability")), likely.get("rationale")]
            )
    add_table("对手反应地图", ["对手", "如果我们", "可能动作", "概率", "理由"], reaction_rows)
    add_table(
        "Payoff 矩阵",
        ["我们的动作", "对方动作", "我方收益", "对方收益", "渠道收益", "基础", "说明"],
        [
            [
                label(row.get("our_action"), action_map),
                label(row.get("opponent_action"), opponent_action_map, text(row.get("opponent_action"))),
                fmt_num(row.get("our_payoff")),
                fmt_num(row.get("opponent_payoff")),
                fmt_num(row.get("channel_payoff")),
                row.get("basis"),
                row.get("notes"),
            ]
            for row in report.get("payoff_matrix", [])
        ],
    )
    add_table(
        "承诺可信度",
        ["动作", "承诺", "可信度", "评分", "说明"],
        [[item.get("action_name"), item.get("commitment"), item.get("level"), fmt_pct(item.get("score")), item.get("notes")] for item in report.get("commitment_tests", [])],
    )
    add_table(
        "动态更新日志",
        ["轮次", "日期", "阶段", "新增信息", "判断", "状态"],
        [[r.get("round"), r.get("date"), r.get("stage"), text(r.get("new_information")), r.get("interim_judgment"), r.get("recommendation_status")] for r in report.get("rounds", [])],
    )
    document.add_heading("风险与边界", level=1)
    for warning in report.get("warnings", []):
        document.add_paragraph(text(warning), style="List Bullet")
    document.save(path)


def write_pdf(path: Path, html: str) -> str | None:
    try:
        from weasyprint import HTML
    except Exception as exc:  # pragma: no cover - fallback depends on environment.
        note = path.with_suffix(".pdf.unavailable.txt")
        note.write_text(f"WeasyPrint is unavailable, so PDF export was skipped: {exc}\nUse the HTML report's Print / Save as PDF action instead.\n", encoding="utf-8")
        return str(note)
    HTML(string=html, base_url=str(ROOT)).write_pdf(path)
    return None


def build_bundle(input_file: Path, output_dir: Path, update_file: Path | None = None, no_pdf: bool = False) -> dict[str, str]:
    request = load_json(input_file)
    if update_file:
        request = merge_update(request, load_json(update_file))
    report = build_report(request)
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = input_file.stem.replace("_", "-")
    canonical_path = output_dir / f"{stem}.canonical.json"
    markdown_path = output_dir / f"{stem}.md"
    html_path = output_dir / f"{stem}.html"
    docx_path = output_dir / f"{stem}.docx"
    pdf_path = output_dir / f"{stem}.pdf"
    write_json(canonical_path, report)
    markdown_path.write_text(render_markdown(report), encoding="utf-8")
    html = render_html(report)
    html_path.write_text(html, encoding="utf-8")
    write_docx(docx_path, report)
    pdf_note = None
    if not no_pdf:
        pdf_note = write_pdf(pdf_path, html)
    return {
        "canonical_json": str(canonical_path),
        "markdown": str(markdown_path),
        "html": str(html_path),
        "docx": str(docx_path),
        "pdf": str(pdf_path if pdf_path.exists() else pdf_note or ""),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a game theory strategy report bundle.")
    parser.add_argument("input_file", help="Structured game theory request JSON.")
    parser.add_argument("output_dir", help="Directory where report artifacts will be written.")
    parser.add_argument("--update", help="Optional update JSON to merge before rendering.")
    parser.add_argument("--no-pdf", action="store_true", help="Skip automated PDF rendering.")
    args = parser.parse_args()
    outputs = build_bundle(
        input_file=Path(args.input_file).resolve(),
        output_dir=Path(args.output_dir).resolve(),
        update_file=Path(args.update).resolve() if args.update else None,
        no_pdf=args.no_pdf,
    )
    print(json.dumps(outputs, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
